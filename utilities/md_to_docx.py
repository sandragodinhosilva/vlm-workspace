#!/usr/bin/env python3
"""
Convert one or more Markdown files to .docx.

INTERPRETER — use this venv (it has python-docx; the others may not):
    /home/sgsilva/vlm-post-training-home-venv/bin/python

Usage (always call with the venv above, NOT a bare `python`):
    /home/sgsilva/vlm-post-training-home-venv/bin/python md_to_docx.py file.md          # file.docx alongside
    /home/sgsilva/vlm-post-training-home-venv/bin/python md_to_docx.py file.md -o out.docx
    /home/sgsilva/vlm-post-training-home-venv/bin/python md_to_docx.py docs/*.md         # batch — sibling .docx each
    /home/sgsilva/vlm-post-training-home-venv/bin/python md_to_docx.py docs/*.md -o out_dir/

Requires: python-docx (already installed in vlm-post-training-home-venv).
If you must use a different venv: pip install python-docx there first.
"""

import argparse
import re
import sys
from pathlib import Path

try:
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
except ImportError:
    sys.exit(
        "python-docx not found in this interpreter.\n"
        "Run with the venv that has it:\n"
        "  /home/sgsilva/vlm-post-training-home-venv/bin/python "
        + " ".join(sys.argv)
    )


# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------

def _strip_inline(text: str) -> str:
    """Remove bold, italic, inline-code, and link markup, keep plain text."""
    text = re.sub(r'\*\*\*(.+?)\*\*\*', r'\1', text)
    text = re.sub(r'\*\*(.+?)\*\*', r'\1', text)
    text = re.sub(r'\*(.+?)\*', r'\1', text)
    text = re.sub(r'`(.+?)`', r'\1', text)
    text = re.sub(r'\[(.+?)\]\(.+?\)', r'\1', text)
    return text


def _is_table_separator(cells: list[str]) -> bool:
    return all(re.match(r'^[-: ]+$', c.strip()) for c in cells if c.strip())


def _flush_table(doc: "Document", rows: list[list[str]]) -> None:
    if not rows:
        return
    n_cols = max(len(r) for r in rows)
    table = doc.add_table(rows=len(rows), cols=n_cols)
    table.style = "Table Grid"
    for i, row in enumerate(rows):
        for j in range(n_cols):
            cell_text = row[j].strip() if j < len(row) else ""
            p = table.cell(i, j).paragraphs[0]
            run = p.add_run(_strip_inline(cell_text))
            run.font.size = Pt(9)
            if i == 0:
                run.bold = True
    doc.add_paragraph()


def _flush_code(doc: "Document", lines: list[str]) -> None:
    if not lines:
        return
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.3)
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run("\n".join(lines))
    run.font.name = "Courier New"
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)


# ---------------------------------------------------------------------------
# Core converter
# ---------------------------------------------------------------------------

def convert(md_path: Path, docx_path: Path) -> None:
    lines = md_path.read_text(encoding="utf-8").splitlines()
    doc = Document()

    # Standard 1-inch margins
    for section in doc.sections:
        for attr in ("left_margin", "right_margin", "top_margin", "bottom_margin"):
            setattr(section, attr, Inches(1))

    in_code = False
    code_lines: list[str] = []
    table_rows: list[list[str]] = []

    for raw in lines:
        line = raw.rstrip()

        # ---- code fence ---------------------------------------------------
        if line.strip().startswith("```"):
            if not in_code:
                in_code = True
                code_lines = []
            else:
                in_code = False
                _flush_code(doc, code_lines)
            continue

        if in_code:
            code_lines.append(line)
            continue

        # ---- table rows ---------------------------------------------------
        if line.strip().startswith("|"):
            cells = line.strip().split("|")[1:-1]  # drop leading/trailing |
            if _is_table_separator(cells):
                continue
            table_rows.append(cells)
            continue
        else:
            _flush_table(doc, table_rows)
            table_rows = []

        # ---- headings -----------------------------------------------------
        m = re.match(r'^(#{1,6})\s+(.*)', line)
        if m:
            level = min(len(m.group(1)), 4)
            doc.add_heading(_strip_inline(m.group(2)), level=level)
            continue

        # ---- horizontal rule ----------------------------------------------
        if re.match(r'^---+$', line.strip()):
            p = doc.add_paragraph("─" * 72)
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after = Pt(2)
            continue

        # ---- unordered list -----------------------------------------------
        m = re.match(r'^(\s*)[-*+] (.*)', line)
        if m:
            indent = len(m.group(1)) // 2  # nesting level
            style = "List Bullet 2" if indent > 0 else "List Bullet"
            doc.add_paragraph(_strip_inline(m.group(2)), style=style)
            continue

        # ---- ordered list -------------------------------------------------
        m = re.match(r'^(\s*)\d+\.\s+(.*)', line)
        if m:
            indent = len(m.group(1)) // 2
            style = "List Number 2" if indent > 0 else "List Number"
            doc.add_paragraph(_strip_inline(m.group(2)), style=style)
            continue

        # ---- blank line ---------------------------------------------------
        if not line.strip():
            continue

        # ---- plain paragraph ---------------------------------------------
        p = doc.add_paragraph(_strip_inline(line.strip()))
        p.paragraph_format.space_after = Pt(4)

    # flush anything left open
    _flush_table(doc, table_rows)
    if in_code:
        _flush_code(doc, code_lines)

    doc.save(str(docx_path))
    print(f"  {md_path} -> {docx_path}")

    # Append a note to the source .md file pointing to the generated .docx
    note_tag = "<!-- md_to_docx -->"
    note_line = f"\n{note_tag}\n> **See also:** [`{docx_path.name}`]({docx_path.name}) — auto-generated Word version of this file.\n"
    content = md_path.read_text(encoding="utf-8")
    # Replace an existing note rather than appending a duplicate
    if note_tag in content:
        content = re.sub(
            rf'{re.escape(note_tag)}\n.*\n',
            note_line.lstrip("\n"),
            content,
        )
    else:
        content = content.rstrip("\n") + note_line
    md_path.write_text(content, encoding="utf-8")


# ---------------------------------------------------------------------------
# CLI
# ---------------------------------------------------------------------------

def main() -> None:
    ap = argparse.ArgumentParser(
        description=__doc__,
        formatter_class=argparse.RawDescriptionHelpFormatter,
    )
    ap.add_argument("inputs", nargs="+", type=Path, help="Markdown file(s) to convert")
    ap.add_argument("-o", "--output", type=Path, default=None,
                    help="Output .docx path (single file) or output directory (batch)")
    args = ap.parse_args()

    inputs = args.inputs
    output = args.output

    if len(inputs) == 1:
        src = inputs[0]
        if output is None:
            dst = src.with_suffix(".docx")
        elif output.is_dir() or str(output).endswith("/"):
            output.mkdir(parents=True, exist_ok=True)
            dst = output / src.with_suffix(".docx").name
        else:
            dst = output
        convert(src, dst)

    else:
        # batch mode
        if output is not None:
            output.mkdir(parents=True, exist_ok=True)
        for src in inputs:
            if not src.exists():
                print(f"  SKIP (not found): {src}", file=sys.stderr)
                continue
            dst = (output / src.with_suffix(".docx").name) if output else src.with_suffix(".docx")
            convert(src, dst)


if __name__ == "__main__":
    main()
